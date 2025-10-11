#!/usr/bin/env python3
"""Check sections in the Word document"""

import sys
sys.path.insert(0, '/Users/paulmarshall/Library/Python/3.9/lib/python/site-packages')

from docx import Document

doc_path = '/Users/paulmarshall/Documents/GitHub/skylerthomas2/kdp/OUT-OF-THE-SWAMP-COMPLETE.docx'

print(f"Opening: {doc_path}")
doc = Document(doc_path)

print(f"\nNumber of sections: {len(doc.sections)}")

for i, section in enumerate(doc.sections):
    print(f"\nSection {i}:")
    print(f"  Start type: {section.start_type}")
    print(f"  Has footer: {len(section.footer.paragraphs) > 0}")

# Find Introduction
for i, para in enumerate(doc.paragraphs):
    if 'Introduction' in para.text and 'Out of the Swamp' in para.text:
        print(f"\nIntroduction found at paragraph {i}: {para.text[:60]}")
        break
    if '\\pagenumbering{arabic}' in para.text:
        print(f"\nMarker found at paragraph {i}: {para.text[:60]}")
