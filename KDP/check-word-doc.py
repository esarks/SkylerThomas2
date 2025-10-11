#!/usr/bin/env python3
"""Check if Word doc has page numbers configured"""

import sys
sys.path.insert(0, '/Users/paulmarshall/Library/Python/3.9/lib/python/site-packages')

from docx import Document

doc_path = '/Users/paulmarshall/Documents/GitHub/SkylerThomas2/KDP/OUT-OF-THE-SWAMP-COMPLETE.docx'

print(f"Checking: {doc_path}")
doc = Document(doc_path)

print(f"\nDocument has {len(doc.sections)} section(s)")

for i, section in enumerate(doc.sections):
    print(f"\nSection {i}:")
    print(f"  Start type: {section.start_type}")
    print(f"  Different first page header/footer: {section.different_first_page_header_footer}")

    footer = section.footer
    print(f"  Footer has {len(footer.paragraphs)} paragraph(s)")

    if footer.paragraphs:
        for j, para in enumerate(footer.paragraphs):
            if para.text.strip():
                print(f"    Footer para {j}: {para.text[:50]}")
            # Check for field codes (page numbers)
            for run in para.runs:
                if hasattr(run._element, 'fldChar'):
                    print(f"    Found field code in para {j}")

print(f"\nTotal paragraphs in document: {len(doc.paragraphs)}")
