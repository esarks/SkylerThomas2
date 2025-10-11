#!/usr/bin/env python3
"""Add section break before Introduction programmatically"""

import sys
sys.path.insert(0, '/Users/paulmarshall/Library/Python/3.9/lib/python/site-packages')

from docx import Document
from docx.enum.section import WD_SECTION

doc_path = '/Users/paulmarshall/Documents/GitHub/skylerthomas2/kdp/OUT-OF-THE-SWAMP-COMPLETE.docx'

print(f"Opening document: {doc_path}")
doc = Document(doc_path)

# Find the paragraph that contains the Introduction heading
intro_para_index = None
for i, para in enumerate(doc.paragraphs):
    if 'Out of the Swamp: How I Found Truth (Introduction)' in para.text:
        intro_para_index = i
        print(f"Found Introduction at paragraph {i}: {para.text[:50]}")
        break
    # Also look for the marker
    if '\\pagenumbering{arabic}' in para.text:
        intro_para_index = i
        print(f"Found page numbering marker at paragraph {i}")
        # Remove the marker text
        para.text = para.text.replace('\\pagenumbering{arabic}', '').strip()
        break

if intro_para_index is None:
    print("ERROR: Could not find Introduction heading or marker")
    sys.exit(1)

# Add section break before this paragraph
if intro_para_index > 0:
    # Insert a new paragraph before the Introduction and add a section break to it
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Get the paragraph where we want to add the section break
    intro_para = doc.paragraphs[intro_para_index]

    # Add a section break by inserting it into the paragraph before the intro
    prev_para = doc.paragraphs[intro_para_index - 1]

    # Create section properties
    sectPr = OxmlElement('w:sectPr')

    # Create a section type element for "next page" break
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)

    # Append section properties to the previous paragraph
    prev_para._element.append(sectPr)

    print(f"Added section break after paragraph {intro_para_index - 1} (before Introduction)")
    print(f"Document now has {len(doc.sections)} section(s)")
else:
    print("ERROR: Introduction is the first paragraph, cannot add section break")
    sys.exit(1)

# Save the document
print("Saving document with section break...")
doc.save(doc_path)
print("✓ Section break added successfully!")
print("\nNow you can run fix-page-numbering.py")
